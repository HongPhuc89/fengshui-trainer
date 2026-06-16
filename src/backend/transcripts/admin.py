import logging
import os
import shutil
import subprocess

from django import forms
from django.contrib import admin, messages
from django.conf import settings
from django.shortcuts import get_object_or_404, redirect
from django.template.response import TemplateResponse
from django.urls import path
from django.utils.html import format_html

from .models import (
    SourceType, TranscriptConfig, TranscriptJob,
    TranscriptApiKey, TranscriptApiKeyUsage,
)
from .tasks import (
    task_download_audio, task_upload_to_gemini,
    task_transcribe_audio, task_translate_transcript,
)

logger = logging.getLogger(__name__)

_ALLOWED_AUDIO_MIME = {
    '.mp3':  'audio/mpeg',
    '.wav':  'audio/wav',
    '.m4a':  'audio/mp4',
    '.ogg':  'audio/ogg',
    '.flac': 'audio/flac',
}
_ALLOWED_AUDIO_EXTS = set(_ALLOWED_AUDIO_MIME.keys())


class TranscriptJobForm(forms.ModelForm):
    class Meta:
        model  = TranscriptJob
        fields = '__all__'

    class Media:
        js = ('admin/transcripts/source_type_toggle.js',)

    def clean(self):
        cleaned = super().clean()
        source_type    = cleaned.get('source_type')
        youtube_url    = (cleaned.get('youtube_url') or '').strip()
        uploaded_audio = cleaned.get('uploaded_audio')

        if source_type == SourceType.YOUTUBE:
            if not youtube_url:
                self.add_error('youtube_url', 'YouTube URL is required for source type YouTube.')
        elif source_type == SourceType.LOCAL_AUDIO:
            is_new = self.instance.pk is None
            if is_new and not uploaded_audio:
                self.add_error('uploaded_audio', 'Please upload an audio file.')
            if uploaded_audio and hasattr(uploaded_audio, 'name'):
                ext = os.path.splitext(uploaded_audio.name)[1].lower()
                if ext not in _ALLOWED_AUDIO_EXTS:
                    self.add_error(
                        'uploaded_audio',
                        f'Unsupported format "{ext}". Allowed: {", ".join(sorted(_ALLOWED_AUDIO_EXTS))}',
                    )
        return cleaned

STATUS_COLORS = {
    'PENDING':      '#aaa',
    'PROCESSING':   '#f90',
    'UPLOADING':    '#ff9800',
    'TRANSCRIBING': '#ff9800',
    'TRANSLATING':  '#ff9800',
    'DONE':         '#4caf50',
    'FAILED':       '#e53935',
    'SKIPPED':      '#2196f3',
}


def _badge(status):
    color = STATUS_COLORS.get(status, '#aaa')
    return format_html(
        '<span style="color:{};font-weight:bold">{}</span>', color, status
    )


@admin.register(TranscriptConfig)
class TranscriptConfigAdmin(admin.ModelAdmin):
    list_display  = ['type', 'model', 'updated_at']
    readonly_fields = ['type', 'updated_at']  # type cannot be changed after creation
    fields = ['type', 'model', 'value', 'updated_at']

    def has_add_permission(self, request):
        return False  # edit only — 2 rows already created via migration

    def has_delete_permission(self, request, obj=None):
        return False  # config rows must not be deleted


class CoverageStatusFilter(admin.SimpleListFilter):
    title = 'Coverage status'
    parameter_name = 'coverage_status'

    def lookups(self, request, model_admin):
        return [
            ('ok',          'OK (≥90%)'),
            ('warn',        'Incomplete (<90%)'),
            ('unverifiable','Unverifiable (ERR)'),
            ('unchecked',   'Not checked yet'),
        ]

    def queryset(self, request, queryset):
        v = self.value()
        if v == 'ok':
            return queryset.filter(transcript_coverage__gte=0.90, transcript_coverage__lte=1.5)
        if v == 'warn':
            return queryset.filter(transcript_coverage__gte=0, transcript_coverage__lt=0.90)
        if v == 'unverifiable':
            # Coverage is None but a warning exists — hallucinated timestamps or parse error
            return queryset.filter(transcript_coverage__isnull=True, step2b_warning__gt='')
        if v == 'unchecked':
            return queryset.filter(transcript_coverage__isnull=True, step2b_warning='')
        return queryset


@admin.register(TranscriptJob)
class TranscriptJobAdmin(admin.ModelAdmin):

    form = TranscriptJobForm
    change_list_template = 'admin/transcripts/transcriptjob_changelist.html'

    list_display = [
        'id', 'title_short', 'source_badge', 'source_url_short',
        'step1_badge', 'step2a_badge', 'step2b_badge', 'step3_badge',
        'overall_badge', 'coverage_badge', 'created_at',
    ]
    list_filter  = ['source_type', 'step1_status', 'step2a_status', 'step2b_status', 'step3_status', CoverageStatusFilter]
    search_fields = ['youtube_url', 'title', 'playlist_url']
    ordering     = ['-created_at']
    actions      = [
        'action_rerun_download', 'action_rerun_upload',
        'action_rerun_transcribe', 'action_rerun_translate',
        'action_trigger_step3',
    ]

    readonly_fields = [
        'uuid', 'title', 'playlist_url', 'audio_file',
        'gemini_file_uri', 'gemini_file_name', 'gemini_uploaded_at',
        'gemini_file_badge', 'gemini_api_key',
        'step1_status', 'step1_error_display', 'step1_timing',
        'step2a_status', 'step2a_error_display', 'step2a_timing',
        'step2b_status', 'step2b_timing', 'step2b_model', 'step2b_error_display',
        'transcript_coverage_display', 'step2b_warning_display',
        'translated_transcript_display',
        'step3_status', 'step3_error_display', 'step3_timing',
        'overall_badge', 'created_at', 'updated_at',
        'rerun_buttons', 'audio_player',
    ]

    fieldsets = [
        ('Job Info', {
            'fields': ['uuid', 'source_type', 'youtube_url', 'uploaded_audio', 'playlist_url', 'title'],
        }),
        ('Re-run Controls', {
            'fields': ['rerun_buttons'],
        }),
        ('Audio Preview', {
            'fields': ['audio_player'],
        }),
        ('Step 1 — Download', {
            'fields': ['step1_status', 'step1_timing', 'audio_file', 'step1_error_display'],
        }),
        ('Step 2a — Upload to Gemini', {
            'fields': [
                'step2a_status', 'step2a_timing', 'step2a_error_display',
                'gemini_api_key',
                'gemini_file_uri', 'gemini_file_name',
                'gemini_uploaded_at', 'gemini_file_badge',
            ],
            'classes': ['collapse'],
        }),
        ('Step 2b — Transcribe (Chinese)', {
            'fields': [
                'step2b_status', 'step2b_model', 'step2b_timing', 'step2b_error_display',
                'transcript_coverage_display', 'step2b_warning_display',
                'raw_transcript',
            ],
            'classes': ['collapse'],
        }),
        ('Step 3 — Translate (Vietnamese)', {
            'fields': ['step3_status', 'step3_timing', 'step3_error_display', 'translated_transcript_display'],
        }),
        ('Metadata', {
            'fields': ['overall_badge', 'created_at', 'updated_at'],
            'classes': ['collapse'],
        }),
    ]

    # --- Custom URLs ---
    def get_urls(self):
        urls = super().get_urls()
        custom = [
            path(
                '<int:job_id>/rerun/<str:step>/',
                self.admin_site.admin_view(self.rerun_step_view),
                name='transcripts_rerun_step',
            ),
            path(
                '<int:job_id>/download-docx/',
                self.admin_site.admin_view(self.download_docx_view),
                name='transcripts_download_docx',
            ),
            path(
                'import-playlist/',
                self.admin_site.admin_view(self.import_playlist_view),
                name='transcripts_import_playlist',
            ),
        ]
        return custom + urls

    def download_docx_view(self, request, job_id):
        import io
        from django.http import HttpResponse
        from docx import Document
        from docx.shared import Pt

        job = get_object_or_404(TranscriptJob, pk=job_id)
        if not job.translated_transcript:
            messages.error(request, f'Job {job_id}: no translated transcript to export.')
            return redirect('admin:transcripts_transcriptjob_change', job_id)

        doc = Document()
        for para in list(doc.paragraphs):
            para._element.getparent().remove(para._element)

        run = doc.add_paragraph().add_run(f'Link clip: {job.youtube_url}')
        run.font.size = Pt(12)
        doc.add_paragraph()
        for line in job.translated_transcript.splitlines():
            doc.add_paragraph().add_run(line).font.size = Pt(12)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        raw_title = (job.title or f'job_{job_id}').strip()
        # Keep Unicode letters/digits and safe punctuation; strip only filesystem-unsafe chars
        safe_title = ''.join(c if c not in r'\/:*?"<>|' else '_' for c in raw_title).strip()
        filename = f'{safe_title[:120]}.docx'

        response = HttpResponse(
            buf.read(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )
        from urllib.parse import quote
        response['Content-Disposition'] = f"attachment; filename*=UTF-8''{quote(filename)}"
        return response

    def rerun_step_view(self, request, job_id, step):
        job = get_object_or_404(TranscriptJob, pk=job_id)  # noqa: F841
        task_map = {
            'step1':  task_download_audio,
            'step2a': task_upload_to_gemini,
            'step2b': task_transcribe_audio,
            'step3':  task_translate_transcript,
        }
        task = task_map.get(step)
        if task:
            task.delay(job_id)
            messages.success(request, f'Job {job_id}: {step} queued.')
        else:
            messages.error(request, f'Invalid step: {step}')
        return redirect('admin:transcripts_transcriptjob_change', job_id)

    def import_playlist_view(self, request):
        """
        GET  → form to enter playlist URL
        POST (step=fetch) → fetch video list from yt-dlp, show checkbox table
        POST (step=confirm) → create TranscriptJob for selected videos
        """
        context = {
            **self.admin_site.each_context(request),
            'title': 'Import from YouTube Playlist',
        }

        if request.method == 'GET':
            return TemplateResponse(request, 'admin/transcripts/import_playlist.html', context)

        step = request.POST.get('step')

        # --- Step fetch: retrieve video list ---
        if step == 'fetch':
            playlist_url = request.POST.get('playlist_url', '').strip()
            if not playlist_url:
                messages.error(request, 'Please enter a playlist URL.')
                return TemplateResponse(request, 'admin/transcripts/import_playlist.html', context)

            try:
                result = subprocess.run(
                    [
                        'yt-dlp',
                        '--flat-playlist',
                        '--print', '%(url)s\t%(title)s',
                        '--no-warnings',
                        playlist_url,
                    ],
                    capture_output=True, text=True, timeout=60, check=True,
                )
                videos = []
                for line in result.stdout.strip().splitlines():
                    parts = line.split('\t', 1)
                    if len(parts) == 2:
                        videos.append({'url': parts[0], 'title': parts[1]})
                    elif len(parts) == 1:
                        videos.append({'url': parts[0], 'title': ''})

                context.update({
                    'playlist_url': playlist_url,
                    'videos': videos,
                    'step': 'confirm',
                })
                return TemplateResponse(request, 'admin/transcripts/import_playlist.html', context)

            except subprocess.TimeoutExpired:
                messages.error(request, 'Fetching playlist timed out. Try again.')
            except subprocess.CalledProcessError as exc:
                messages.error(request, f'yt-dlp error: {exc.stderr[:500]}')
            except Exception as exc:
                messages.error(request, str(exc))

            return TemplateResponse(request, 'admin/transcripts/import_playlist.html', context)

        # --- Step confirm: create jobs ---
        if step == 'confirm':
            playlist_url   = request.POST.get('playlist_url', '')
            selected_urls  = request.POST.getlist('selected_videos')
            selected_titles = {
                url: request.POST.get(f'title_{url}', '')
                for url in selected_urls
            }

            if not selected_urls:
                messages.warning(request, 'No videos selected.')
                return redirect('admin:transcripts_import_playlist')

            # Check for existing jobs to avoid duplicates
            existing_urls = set(
                TranscriptJob.objects.filter(youtube_url__in=selected_urls)
                .values_list('youtube_url', flat=True)
            )

            created = 0
            skipped = 0
            for url in selected_urls:
                if url in existing_urls:
                    skipped += 1
                    continue
                job = TranscriptJob.objects.create(
                    youtube_url=url,
                    playlist_url=playlist_url,
                    title=selected_titles.get(url, '')[:500],
                )
                pipeline = (
                    task_download_audio.si(job.pk)
                    | task_upload_to_gemini.si(job.pk)
                    | task_transcribe_audio.si(job.pk)
                    | task_translate_transcript.si(job.pk)
                )
                pipeline.delay()
                created += 1

            if skipped:
                messages.warning(request, f'{skipped} video(s) skipped — job already exists.')
            messages.success(request, f'{created} job(s) created and queued.')
            return redirect('admin:transcripts_transcriptjob_changelist')

        return TemplateResponse(request, 'admin/transcripts/import_playlist.html', context)

    def _audio_mime(self, obj):
        ext = os.path.splitext(str(obj.audio_file))[1].lower()
        return _ALLOWED_AUDIO_MIME.get(ext, 'audio/mpeg')

    # --- Audio Player ---
    @admin.display(description='Audio Preview')
    def audio_player(self, obj):
        mode = obj.audio_serve_mode
        if mode == 'gemini':
            # Gemini file URI cannot be streamed directly in browser
            # Show local audio player if file exists, fallback to info message
            if obj.audio_file:
                url  = f'{settings.MEDIA_URL}{obj.audio_file}'
                mime = self._audio_mime(obj)
                return format_html(
                    '<audio controls style="width:100%">'
                    '<source src="{}" type="{}">'
                    '</audio>'
                    '<small style="color:#aaa">Gemini file valid until 48h • Serving from local</small>',
                    url, mime,
                )
            return format_html(
                '<small style="color:#2196f3">Gemini file URI valid (< 48h) — local file still available</small>'
            )
        if mode == 'local':
            url  = f'{settings.MEDIA_URL}{obj.audio_file}'
            mime = self._audio_mime(obj)
            return format_html(
                '<audio controls style="width:100%">'
                '<source src="{}" type="{}">'
                '</audio>',
                url, mime,
            )
        return format_html('<small style="color:#aaa">Audio file expired (> 15 days) — deleted</small>')

    # --- Gemini file badge ---
    @admin.display(description='Gemini File Status')
    def gemini_file_badge(self, obj):
        if not obj.gemini_file_uri:
            return '—'
        if obj.gemini_file_valid:
            return format_html('<span style="color:#4caf50;font-weight:bold">✓ Valid (< 48h)</span>')
        return format_html('<span style="color:#aaa">✗ Expired (> 48h) — re-run Step 2a to re-upload</span>')

    # --- List display helpers ---
    @admin.display(description='Title')
    def title_short(self, obj):
        return (obj.title or '(no title)')[:50]

    @admin.display(description='Source')
    def source_badge(self, obj):
        if obj.source_type == SourceType.LOCAL_AUDIO:
            return format_html('<span style="color:#2196f3;font-weight:bold">LOCAL</span>')
        return format_html('<span style="color:#4caf50;font-weight:bold">YT</span>')

    @admin.display(description='URL / File')
    def source_url_short(self, obj):
        if obj.source_type == SourceType.LOCAL_AUDIO:
            name = os.path.basename(str(obj.audio_file)) if obj.audio_file else '—'
            return name[:45]
        if obj.youtube_url:
            return format_html(
                '<a href="{}" target="_blank">{}</a>',
                obj.youtube_url, obj.youtube_url[:45],
            )
        return '—'

    @admin.display(description='S1')
    def step1_badge(self, obj): return _badge(obj.step1_status)

    @admin.display(description='S2a')
    def step2a_badge(self, obj): return _badge(obj.step2a_status)

    @admin.display(description='S2b')
    def step2b_badge(self, obj): return _badge(obj.step2b_status)

    @admin.display(description='S3')
    def step3_badge(self, obj): return _badge(obj.step3_status)

    @admin.display(description='Overall')
    def overall_badge(self, obj): return _badge(obj.overall_status)

    def _step_timing(self, started_at, finished_at):
        if not started_at:
            return '—'
        start_str = started_at.strftime('%Y-%m-%d %H:%M:%S')
        if not finished_at:
            return format_html('<span style="color:#f90">{} → running…</span>', start_str)
        duration = finished_at - started_at
        total_secs = int(duration.total_seconds())
        if total_secs < 60:
            dur_str = f'{total_secs}s'
        else:
            dur_str = f'{total_secs // 60}m {total_secs % 60}s'
        finish_str = finished_at.strftime('%H:%M:%S')
        return format_html(
            '{} → {} <span style="color:#aaa">({})</span>',
            start_str, finish_str, dur_str,
        )

    @admin.display(description='Timing')
    def step1_timing(self, obj):
        return self._step_timing(obj.step1_started_at, obj.step1_finished_at)

    @admin.display(description='Timing')
    def step2a_timing(self, obj):
        return self._step_timing(obj.step2a_started_at, obj.step2a_finished_at)

    @admin.display(description='Timing')
    def step2b_timing(self, obj):
        return self._step_timing(obj.step2b_started_at, obj.step2b_finished_at)

    @admin.display(description='Timing')
    def step3_timing(self, obj):
        return self._step_timing(obj.step3_started_at, obj.step3_finished_at)

    @admin.display(description='Coverage')
    def coverage_badge(self, obj):
        if obj.transcript_coverage is None:
            return format_html('<span style="color:#aaa">—</span>')
        pct = obj.transcript_coverage * 100
        color = '#4caf50' if pct >= 90 else '#e53935'
        return format_html(
            '<span style="color:{};font-weight:bold">{}</span>',
            color, f'{pct:.1f}%',
        )

    @admin.display(description='Transcript Coverage')
    def transcript_coverage_display(self, obj):
        if obj.transcript_coverage is None:
            return '—'
        pct = obj.transcript_coverage * 100
        color = '#4caf50' if pct >= 90 else '#e53935'
        return format_html(
            '<span style="color:{};font-weight:bold;font-size:16px">{}</span>'
            ' <span style="color:#aaa;font-size:12px">(≥90% = OK, <90% = incomplete)</span>',
            color, f'{pct:.1f}%',
        )

    @admin.display(description='Step 2b Warning')
    def step2b_warning_display(self, obj):
        if not obj.step2b_warning:
            return '—'
        return format_html(
            '<span style="color:#f90;font-weight:bold">⚠ {}</span>',
            obj.step2b_warning,
        )

    # --- Error displays ---
    def _error_field(self, error_text):
        if not error_text:
            return '—'
        return format_html(
            '<pre style="white-space:pre-wrap;color:#e53935;font-size:12px">{}</pre>',
            error_text,
        )

    @admin.display(description='Step 1 Error')
    def step1_error_display(self, obj): return self._error_field(obj.step1_error)

    @admin.display(description='Step 2a Error')
    def step2a_error_display(self, obj): return self._error_field(obj.step2a_error)

    @admin.display(description='Step 2b Error')
    def step2b_error_display(self, obj): return self._error_field(obj.step2b_error)

    @admin.display(description='Step 3 Error')
    def step3_error_display(self, obj): return self._error_field(obj.step3_error)

    # --- Transcript displays ---
    def _transcript_field(self, text, max_height='400px'):
        if not text:
            return '—'
        return format_html(
            '<pre style="white-space:pre-wrap;max-height:{};overflow-y:auto;'
            'font-family:monospace;font-size:13px;line-height:1.6">{}</pre>',
            max_height, text,
        )

    @admin.display(description='Translated Transcript (Vietnamese)')
    def translated_transcript_display(self, obj):
        if not obj.translated_transcript:
            return '—'
        return format_html(
            '<div style="margin-bottom:8px;display:flex;gap:8px;align-items:center">'
            '<button type="button" onclick="'
            'navigator.clipboard.writeText(document.getElementById(\'trans-text-{}\').innerText)'
            '.then(function(){{this.textContent=\'✓ Copied!\';setTimeout(()=>this.textContent=\'Copy\',2000)}}.bind(this))'
            '" style="padding:4px 12px;cursor:pointer">Copy</button>'
            '<a class="button" href="/admin/transcripts/transcriptjob/{}/download-docx/" '
            'style="padding:4px 12px;text-decoration:none">⬇ Download Word</a>'
            '</div>'
            '<pre id="trans-text-{}" style="white-space:pre-wrap;max-height:600px;overflow-y:auto;'
            'font-family:monospace;font-size:13px;line-height:1.6">{}</pre>',
            obj.pk, obj.pk, obj.pk, obj.translated_transcript,
        )

    # --- Re-run buttons on detail page ---
    @admin.display(description='Re-run Controls')
    def rerun_buttons(self, obj):
        if not obj.pk:
            return '(save job first)'
        base = f'/admin/transcripts/transcriptjob/{obj.pk}/rerun'
        if obj.source_type == SourceType.LOCAL_AUDIO:
            return format_html(
                '<a class="button" href="{}/step2a/" style="margin:4px">▶ Step 2a (Upload)</a>'
                '<a class="button" href="{}/step2b/" style="margin:4px">▶ Step 2b (Transcribe)</a>'
                '<a class="button" href="{}/step3/" style="margin:4px">▶ Step 3 (Translate)</a>',
                base, base, base,
            )
        return format_html(
            '<a class="button" href="{}/step1/" style="margin:4px">▶ Step 1 (Download)</a>'
            '<a class="button" href="{}/step2a/" style="margin:4px">▶ Step 2a (Upload)</a>'
            '<a class="button" href="{}/step2b/" style="margin:4px">▶ Step 2b (Transcribe)</a>'
            '<a class="button" href="{}/step3/" style="margin:4px">▶ Step 3 (Translate)</a>',
            base, base, base, base,
        )

    # --- Bulk actions ---
    @admin.action(description='▶ Re-run Step 1 (Download)')
    def action_rerun_download(self, request, queryset):
        queued  = 0
        skipped = 0
        for job in queryset:
            if job.source_type == SourceType.LOCAL_AUDIO:
                skipped += 1
                continue
            task_download_audio.delay(job.pk)
            queued += 1
        if queued:
            self.message_user(request, f'{queued} download task(s) queued.')
        if skipped:
            self.message_user(
                request,
                f'{skipped} LOCAL_AUDIO job(s) skipped — Step 1 (Download) does not apply.',
                level=messages.WARNING,
            )

    @admin.action(description='▶ Re-run Step 2a (Upload to Gemini)')
    def action_rerun_upload(self, request, queryset):
        for job in queryset:
            task_upload_to_gemini.delay(job.pk)
        self.message_user(request, f'{queryset.count()} upload task(s) queued.')

    @admin.action(description='▶ Re-run Step 2b (Transcribe)')
    def action_rerun_transcribe(self, request, queryset):
        for job in queryset:
            task_transcribe_audio.delay(job.pk)
        self.message_user(request, f'{queryset.count()} transcribe task(s) queued.')

    @admin.action(description='▶ Re-run Step 3 (Translate)')
    def action_rerun_translate(self, request, queryset):
        for job in queryset:
            task_translate_transcript.delay(job.pk)
        self.message_user(request, f'{queryset.count()} translate task(s) queued.')

    @admin.action(description='▶ Trigger Step 3 — translate raw_transcript (for manually-entered transcripts)')
    def action_trigger_step3(self, request, queryset):
        from .models import StepStatus
        queued = 0
        skipped = 0
        for job in queryset:
            if not job.raw_transcript:
                skipped += 1
                continue
            job.step2b_status = StepStatus.DONE
            job.save(update_fields=['step2b_status', 'updated_at'])
            task_translate_transcript.apply_async(args=[job.pk])
            queued += 1
        if queued:
            self.message_user(request, f'{queued} job(s) queued for Step 3 translation.')
        if skipped:
            self.message_user(
                request,
                f'{skipped} job(s) skipped — raw_transcript is empty.',
                level=messages.WARNING,
            )

    def _setup_local_audio_job(self, request, job, uploaded_file):
        """
        Move the temp-uploaded file to transcripts/<pk>/audio.<ext>,
        set audio_file and step1_status=SKIPPED on the job.
        """
        from .models import StepStatus
        original_name = uploaded_file.name if hasattr(uploaded_file, 'name') else ''
        ext = os.path.splitext(original_name)[1].lower() or '.audio'
        src = os.path.join(settings.MEDIA_ROOT, str(job.uploaded_audio))
        dest_dir  = os.path.join(settings.MEDIA_ROOT, 'transcripts', str(job.pk))
        dest_name = f'audio{ext}'
        dest      = os.path.join(dest_dir, dest_name)
        os.makedirs(dest_dir, exist_ok=True)
        if not os.path.exists(src):
            messages.error(request, f'Job {job.pk}: uploaded file not found at {src}.')
            return
        try:
            shutil.move(src, dest)
        except (OSError, shutil.Error) as exc:
            messages.error(request, f'Job {job.pk}: could not move audio file — {exc}')
            return
        try:
            os.rmdir(os.path.dirname(src))
        except OSError:
            pass

        job.audio_file   = f'transcripts/{job.pk}/{dest_name}'
        job.step1_status = StepStatus.SKIPPED
        if not job.title and original_name:
            job.title = os.path.splitext(original_name)[0][:500]
        job.save(update_fields=['audio_file', 'step1_status', 'title', 'updated_at'])

    # --- Auto-start pipeline on create ---
    def save_model(self, request, obj, form, change):
        from .models import StepStatus
        is_new = obj.pk is None
        if is_new and obj.source_type == SourceType.LOCAL_AUDIO:
            # Mark step1 SKIPPED before initial save so overall_status is consistent
            obj.step1_status = StepStatus.SKIPPED
        super().save_model(request, obj, form, change)
        if is_new:
            if obj.source_type == SourceType.LOCAL_AUDIO:
                uploaded_file = form.cleaned_data.get('uploaded_audio')
                if uploaded_file:
                    self._setup_local_audio_job(request, obj, uploaded_file)
                pipeline = (
                    task_upload_to_gemini.si(obj.pk)
                    | task_transcribe_audio.si(obj.pk)
                    | task_translate_transcript.si(obj.pk)
                )
                pipeline.delay()
                messages.info(request, f'Job {obj.pk}: LOCAL_AUDIO pipeline (Steps 2a→3) queued.')
            else:
                pipeline = (
                    task_download_audio.si(obj.pk)
                    | task_upload_to_gemini.si(obj.pk)
                    | task_transcribe_audio.si(obj.pk)
                    | task_translate_transcript.si(obj.pk)
                )
                pipeline.delay()
                messages.info(request, f'Job {obj.pk}: Full pipeline queued.')

    # --- Cleanup audio file on delete ---
    def _delete_audio_file(self, job):
        if job.audio_file:
            path = job.audio_file_path
            if path and os.path.exists(path):
                os.remove(path)
                try:
                    os.rmdir(os.path.dirname(path))
                except OSError:
                    pass

    def delete_model(self, request, obj):
        self._delete_audio_file(obj)
        super().delete_model(request, obj)

    def delete_queryset(self, request, queryset):
        for job in queryset:
            self._delete_audio_file(job)
        super().delete_queryset(request, queryset)


# --- API Key Pool ---

class TranscriptApiKeyUsageInline(admin.TabularInline):
    model           = TranscriptApiKeyUsage
    extra           = 0
    can_delete      = False
    fields          = ['model_name', 'rpd_count', 'rpd_reset_at', 'exhausted_until', 'usage_status']
    readonly_fields = ['model_name', 'rpd_count', 'rpd_reset_at', 'exhausted_until', 'usage_status']

    @admin.display(description='Status')
    def usage_status(self, obj):
        from django.utils import timezone
        from django.conf import settings as dj_settings
        if obj.exhausted_until and obj.exhausted_until > timezone.now():
            until = obj.exhausted_until.strftime('%d/%m %H:%M')
            return format_html('<span style="color:#e53935;font-weight:bold">⛔ Until {}</span>', until)
        rpd_limit = dj_settings.GEMINI_RPD_LIMIT
        if obj.rpd_count >= rpd_limit:
            return format_html(
                '<span style="color:#f90;font-weight:bold">⚠ RPD limit ({}/{})</span>',
                obj.rpd_count, rpd_limit,
            )
        return format_html(
            '<span style="color:#4caf50;font-weight:bold">✓ {}/{}</span>',
            obj.rpd_count, rpd_limit,
        )


@admin.register(TranscriptApiKey)
class TranscriptApiKeyAdmin(admin.ModelAdmin):
    list_display    = [
        'label', 'is_active', 'request_count', 'last_used_at',
        'usage_flash25', 'usage_flash30_lite', 'usage_flash30', 'usage_flash30_prv', 'usage_file_api',
        'key_status',
    ]
    list_editable   = ['is_active']
    search_fields   = ['label']
    ordering        = ['label']
    inlines         = [TranscriptApiKeyUsageInline]
    actions         = ['action_reset_all_exhausted']

    readonly_fields = ['api_key_masked', 'request_count', 'last_used_at', 'created_at']
    fields          = ['label', 'api_key', 'api_key_masked', 'is_active',
                       'request_count', 'last_used_at', 'created_at']

    def get_queryset(self, request):
        return super().get_queryset(request).prefetch_related('usages')

    def _usage_badge(self, obj, model_name):
        from django.utils import timezone
        from django.conf import settings as dj_settings
        usage = next((u for u in obj.usages.all() if u.model_name == model_name), None)
        if usage is None:
            return format_html('<span style="color:#aaa">—</span>')
        rpd_limit = dj_settings.GEMINI_RPD_LIMIT
        if usage.exhausted_until and usage.exhausted_until > timezone.now():
            until = usage.exhausted_until.strftime('%H:%M')
            return format_html(
                '<span style="color:#e53935;font-weight:bold">⛔ {}</span>', until,
            )
        if usage.rpd_count >= rpd_limit:
            return format_html(
                '<span style="color:#f90;font-weight:bold">⚠ {}/{}</span>',
                usage.rpd_count, rpd_limit,
            )
        return format_html(
            '<span style="color:#4caf50;font-weight:bold">✓ {}/{}</span>',
            usage.rpd_count, rpd_limit,
        )

    @admin.display(description='2.5-flash')
    def usage_flash25(self, obj):
        return self._usage_badge(obj, 'gemini-2.5-flash')

    @admin.display(description='3.0-flash')
    def usage_flash30(self, obj):
        return self._usage_badge(obj, 'gemini-3.0-flash')

    @admin.display(description='3.0-lite')
    def usage_flash30_lite(self, obj):
        return self._usage_badge(obj, 'gemini-3.0-flash-lite')

    @admin.display(description='3-flash-prv')
    def usage_flash30_prv(self, obj):
        return self._usage_badge(obj, 'gemini-3-flash-preview')

    @admin.display(description='File API')
    def usage_file_api(self, obj):
        return self._usage_badge(obj, 'gemini-file-api')

    @admin.display(description='Status')
    def key_status(self, obj):
        if not obj.is_active:
            return format_html('<span style="color:#aaa">— Disabled</span>')
        return format_html('<span style="color:#4caf50;font-weight:bold">✓ Active</span>')

    @admin.display(description='API Key (masked)')
    def api_key_masked(self, obj):
        if not obj.api_key:
            return '—'
        return f'****{obj.api_key[-8:]}'

    @admin.action(description='Reset exhausted_until for all usages of selected keys')
    def action_reset_all_exhausted(self, request, queryset):
        updated = TranscriptApiKeyUsage.objects.filter(
            api_key__in=queryset,
        ).update(exhausted_until=None)
        self.message_user(request, f'{updated} usage row(s) quota reset.')
