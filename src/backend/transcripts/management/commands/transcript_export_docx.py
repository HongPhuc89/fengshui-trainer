"""
Management command: transcript_export_docx

Export the translated transcript of every TranscriptJob (step3 DONE) to a
DOCX file.  One file per job; filename is derived from the job title.

File structure:
  Line 1 : Link clip: <youtube_url or playlist_url>
  Line 2 : (blank)
  Line 3+: translated transcript text

Usage:
    docker-compose -f docker/docker-compose.yml exec web \\
        python manage.py transcript_export_docx

Options:
    --playlist-url URL    Filter jobs by playlist_url (exports only that playlist)
    --output-dir DIR      Directory to write .docx files
                          (default: media/exports or media/exports/<playlist-slug>)
    --job-ids IDS         Comma-separated job IDs to export (default: all DONE jobs)
    --exclude-ids IDS     Comma-separated job IDs to skip (applied when --job-ids is not set)
    --force               Overwrite existing files (default: skip if file already exists)
"""

import os
import re

from django.conf import settings
from django.core.management.base import BaseCommand


def _sanitize_filename(name: str) -> str:
    """Strip characters that are unsafe in filenames, collapse whitespace."""
    name = re.sub(r'[\\/:*?"<>|]', '', name)
    name = re.sub(r'\s+', ' ', name).strip()
    return name or 'untitled'


class Command(BaseCommand):
    help = 'Export translated transcripts to DOCX files (one file per job)'

    def add_arguments(self, parser):
        parser.add_argument(
            '--playlist-url',
            default='',
            metavar='URL',
            help='Filter jobs by playlist_url',
        )
        parser.add_argument(
            '--output-dir',
            default='',
            metavar='DIR',
            help='Directory to write .docx files (default: MEDIA_ROOT/exports or MEDIA_ROOT/exports/<playlist-slug>)',
        )
        parser.add_argument(
            '--job-ids',
            default='',
            metavar='IDS',
            help='Comma-separated job IDs to export (default: all jobs with step3=DONE)',
        )
        parser.add_argument(
            '--exclude-ids',
            default='',
            metavar='IDS',
            help='Comma-separated job IDs to skip (only applied when --job-ids is not set)',
        )
        parser.add_argument(
            '--force',
            action='store_true',
            help='Overwrite existing .docx files (default: skip)',
        )

    def handle(self, *args, **options):
        try:
            from docx import Document
            from docx.shared import Pt
        except ImportError:
            self.stderr.write(
                self.style.ERROR(
                    'python-docx is not installed. '
                    'Run: pip install python-docx'
                )
            )
            return

        from transcripts.models import TranscriptJob, StepStatus

        playlist_url    = options['playlist_url'].strip()
        job_ids_raw     = options['job_ids'].strip()
        exclude_ids_raw = options['exclude_ids'].strip()

        # Resolve output dir
        if options['output_dir']:
            output_dir = options['output_dir']
        elif playlist_url:
            slug = _sanitize_filename(playlist_url.split('list=')[-1] if 'list=' in playlist_url else playlist_url)
            output_dir = os.path.join(settings.MEDIA_ROOT, 'exports', slug[:60])
        else:
            output_dir = os.path.join(settings.MEDIA_ROOT, 'exports')
        os.makedirs(output_dir, exist_ok=True)

        # Resolve job queryset
        if job_ids_raw:
            ids = [int(i.strip()) for i in job_ids_raw.split(',') if i.strip()]
            jobs = list(TranscriptJob.objects.filter(pk__in=ids).order_by('pk'))
            if not jobs:
                self.stdout.write(self.style.WARNING('No jobs found for the given IDs.'))
                return
        else:
            qs = TranscriptJob.objects.filter(step3_status=StepStatus.DONE)
            if playlist_url:
                qs = qs.filter(playlist_url=playlist_url)
            if exclude_ids_raw:
                exclude_ids = [int(i.strip()) for i in exclude_ids_raw.split(',') if i.strip()]
                qs = qs.exclude(pk__in=exclude_ids)
            jobs = list(qs.order_by('title', 'pk'))
            if not jobs:
                msg = f'No jobs with step3=DONE found'
                msg += f' for playlist: {playlist_url}' if playlist_url else ''
                self.stdout.write(self.style.WARNING(msg + '.'))
                return

        self.stdout.write(f'Exporting {len(jobs)} job(s) to: {output_dir}')
        self.stdout.write('')

        exported = 0
        skipped = 0
        failed = 0

        for job in jobs:
            label = f'Job #{job.pk} — {(job.title or job.youtube_url)[:60]}'

            if not job.translated_transcript:
                self.stdout.write(
                    self.style.WARNING(f'{label}')
                )
                self.stdout.write('  skipped: translated_transcript is empty')
                skipped += 1
                continue

            safe_title = _sanitize_filename(job.title or f'job_{job.pk}')
            filename = f'{safe_title}.docx'
            filepath = os.path.join(output_dir, filename)

            if os.path.exists(filepath) and not options['force']:
                self.stdout.write(f'{label}')
                self.stdout.write(f'  skipped: {filename} already exists (use --force to overwrite)')
                skipped += 1
                continue

            try:
                doc = Document()

                # Remove default empty paragraph added by python-docx
                for para in list(doc.paragraphs):
                    p = para._element
                    p.getparent().remove(p)

                # Line 1: clip URL (youtube_url for YT jobs, playlist_url for LOCAL_AUDIO)
                source_ref = job.youtube_url or job.playlist_url or f'job #{job.pk}'
                url_para = doc.add_paragraph()
                run = url_para.add_run(f'Link clip: {source_ref}')
                run.font.size = Pt(12)

                # Line 2: blank
                doc.add_paragraph()

                # Lines 3+: translated transcript (preserve line breaks as separate paragraphs)
                for line in job.translated_transcript.splitlines():
                    para = doc.add_paragraph()
                    run = para.add_run(line)
                    run.font.size = Pt(12)

                doc.save(filepath)
                self.stdout.write(self.style.SUCCESS(f'{label}'))
                self.stdout.write(f'  → {filepath}')
                exported += 1

            except Exception as exc:
                self.stdout.write(self.style.ERROR(f'{label}'))
                self.stdout.write(f'  FAILED: {exc}')
                failed += 1

        self.stdout.write('')
        self.stdout.write('─' * 60)
        self.stdout.write(f'Summary: {exported} exported, {skipped} skipped, {failed} failed')
